from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_center(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def main():
    doc = Document()

    # Portada
    add_title(doc, "Fundamentos de Economia y Finanzas")
    add_center(doc, "Caso practico integral", size=14, bold=True)
    add_center(doc, "Primer reto profesional: Finanzas entre temporadas", size=12)
    doc.add_paragraph("")
    add_center(doc, "Propuesta para: Director General de Temporadas del Sur", size=11)
    add_center(doc, "Elaborado por: [Tu nombre completo]", size=11)
    add_center(doc, "Curso: Fundamentos de Economia y Finanzas", size=11)
    add_center(doc, f"Fecha: {date.today().strftime('%d/%m/%Y')}", size=11)
    doc.add_page_break()

    add_heading(doc, "1) Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "Temporadas del Sur vende la mayor parte de sus productos en agosto y septiembre, "
        "y depende mucho del mercado de Estados Unidos. Esto hace que la empresa tenga meses "
        "con poca liquidez y un riesgo alto si cambian los aranceles o el tipo de cambio. "
        "La propuesta busca mejorar el flujo de efectivo, reducir riesgos y abrir nuevos mercados."
    )
    doc.add_paragraph(
        "La idea principal es combinar acciones de corto plazo (ordenar caja, cobros, pagos "
        "y cobertura cambiaria) con acciones de largo plazo (diversificacion de mercados, "
        "monitoreo economico y mejor uso del sistema financiero)."
    )

    add_heading(doc, "2) Situacion y contexto economico (Apartado A)", level=1)
    doc.add_paragraph(
        "Con base en el caso, la empresa tiene ventas anuales estimadas de 3.6 millones de pesos. "
        "El 70% se va a costos de produccion y 10% a gastos administrativos. Esto deja un margen "
        "reducido cuando hay presion por inflacion o por tipo de cambio."
    )
    doc.add_paragraph("Conceptos clave aplicados en lenguaje sencillo:")
    add_bullet(doc, "Oferta y demanda: la demanda sube en meses frios y baja el resto del anio.")
    add_bullet(
        doc,
        "Elasticidad: en algunos mercados los clientes son sensibles al precio; subir precio puede bajar ventas.",
    )
    add_bullet(
        doc,
        "Inflacion: sube costos (hilo, transporte, energia, sueldos) y obliga a revisar precios.",
    )
    add_bullet(
        doc,
        "Tasa de interes: si sube, el credito bancario se vuelve mas caro para financiar inventario.",
    )
    add_bullet(
        doc,
        "Tipo de cambio peso-dolar: afecta ingresos por exportacion y costos ligados al dolar.",
    )
    add_bullet(
        doc,
        "Instituciones financieras: bancos, fintech y organismos de apoyo pueden dar credito, coberturas e inversion.",
    )
    doc.add_paragraph(
        "En resumen, la empresa necesita planear con datos macroeconomicos para no reaccionar tarde "
        "cuando cambien las condiciones del mercado."
    )

    add_heading(doc, "3) Propuestas de corto plazo (Apartado C)", level=1)
    doc.add_paragraph("3.1 Mejora del flujo de caja (0 a 6 meses)")
    add_bullet(
        doc,
        "Hacer un presupuesto semanal de caja para saber faltantes y sobrantes con anticipacion.",
    )
    add_bullet(
        doc,
        "Negociar con proveedores ampliar credito de 45 a 60 dias en temporada baja.",
    )
    add_bullet(
        doc,
        "Solicitar una linea de credito revolvente para capital de trabajo y usarla solo cuando sea necesario.",
    )
    add_bullet(
        doc,
        "Reducir el uso de efectivo y mover operaciones a banca digital para seguridad y control.",
    )

    doc.add_paragraph("3.2 Reestructuracion de cobros y pagos")
    add_bullet(doc, "Cobrar anticipos del 30% en pedidos grandes de exportacion.")
    add_bullet(doc, "Ofrecer descuento pequeno por pronto pago a clientes confiables.")
    add_bullet(
        doc,
        "Programar pagos de gastos no urgentes en semanas de mayor entrada de efectivo.",
    )

    doc.add_paragraph("3.3 Cobertura de riesgos (tipo de cambio y diesel)")
    add_bullet(
        doc,
        "Usar forwards de tipo de cambio para asegurar el valor en pesos de ventas en dolares.",
    )
    add_bullet(
        doc,
        "Explorar coberturas de combustible para reducir sorpresas en el costo del diesel.",
    )
    add_bullet(
        doc,
        "Definir una politica simple: cubrir entre 50% y 70% de exposicion de los siguientes 3 a 6 meses.",
    )

    add_heading(doc, "4) Propuestas de largo plazo (Apartado B)", level=1)
    doc.add_paragraph("4.1 Sistema de monitoreo economico")
    add_bullet(doc, "Crear tablero mensual con: inflacion, tipo de cambio, tasa de interes, PIB y empleo.")
    add_bullet(doc, "Asignar responsable y fecha de revision cada mes.")
    add_bullet(doc, "Definir alertas (por ejemplo: si el dolar sube mas de 5%, activar cobertura adicional).")

    doc.add_paragraph("4.2 Diversificacion de mercados (Europa y America Latina)")
    add_bullet(
        doc,
        "Prioridad 1: Colombia y Espana por idioma, canales comerciales y menor barrera de entrada.",
    )
    add_bullet(
        doc,
        "Prioridad 2: Brasil y Alemania como crecimiento gradual con analisis logistico y regulatorio.",
    )
    add_bullet(
        doc,
        "Objetivo: bajar dependencia de EE. UU. de 85% a 65% en 3 anios.",
    )
    add_bullet(
        doc,
        "Hacer pruebas de precio por mercado para estimar elasticidad y ajustar estrategia comercial.",
    )

    doc.add_paragraph("4.3 Uso de inversiones y ahorro de corto plazo")
    add_bullet(
        doc,
        "Colocar excedentes temporales en instrumentos de bajo riesgo y alta liquidez (Cetes o fondos de deuda de corto plazo).",
    )
    add_bullet(doc, "Separar un fondo de contingencia de al menos 2 meses de gastos fijos.")

    add_heading(doc, "5) Analisis de mercado y elasticidad (reto 3 y 4)", level=1)
    doc.add_paragraph(
        "Para saber si el producto es elastico o inelastico, se recomienda hacer pruebas controladas "
        "de precio por canal y por mercado durante periodos cortos."
    )
    doc.add_paragraph("Metodo simple propuesto:")
    add_bullet(doc, "Seleccionar 2 o 3 ciudades por mercado.")
    add_bullet(doc, "Aplicar variaciones de precio pequenas (+/-5%) por 4 a 6 semanas.")
    add_bullet(doc, "Medir cambio porcentual en cantidad vendida.")
    add_bullet(doc, "Calcular elasticidad-precio = % cambio en cantidad / % cambio en precio.")
    add_bullet(
        doc,
        "Si |elasticidad| > 1, la demanda es elastica (conviene cuidar aumentos de precio).",
    )
    add_bullet(
        doc,
        "Si |elasticidad| < 1, la demanda es inelastica (hay mas espacio para ajustar precios).",
    )

    doc.add_paragraph(
        "Con esta informacion, la empresa puede aproximar un precio de equilibrio por mercado para "
        "maximizar ingresos sin perder demasiadas ventas."
    )

    add_heading(doc, "6) Simulacion financiera simple (apoyo a decisiones)", level=1)
    doc.add_paragraph("Escenario base anual (segun caso):")
    add_bullet(doc, "Ventas: 3.6 millones MXN")
    add_bullet(doc, "Costos de produccion (70%): 2.52 millones MXN")
    add_bullet(doc, "Gastos administrativos (10%): 0.36 millones MXN")
    add_bullet(doc, "Margen antes de gastos financieros e impuestos: 0.72 millones MXN")

    doc.add_paragraph("Escenario estresado (inflacion +3% y menor demanda externa):")
    add_bullet(doc, "Aumento de costos y presion en flujo de efectivo.")
    add_bullet(doc, "Mayor necesidad de credito de corto plazo en meses de baja venta.")
    add_bullet(doc, "Importancia critica de coberturas y diversificacion de mercados.")

    add_heading(doc, "7) Comparacion de resultados y metrica de control", level=1)
    doc.add_paragraph(
        "Para verificar si la estrategia funciona, se recomienda comparar cada mes:"
    )
    add_bullet(doc, "Ventas reales vs presupuesto.")
    add_bullet(doc, "Margen bruto real vs margen objetivo.")
    add_bullet(doc, "Desviacion del flujo de caja.")
    add_bullet(doc, "Porcentaje de ventas cubiertas contra tipo de cambio.")
    add_bullet(doc, "Participacion de mercados distintos a EE. UU.")
    doc.add_paragraph(
        "Una metrica de error sencilla es el Error Porcentual Medio Absoluto (MAPE) entre "
        "ventas proyectadas y ventas reales por mes."
    )

    add_heading(doc, "8) Conclusiones (respuesta directa)", level=1)
    doc.add_paragraph(
        "Si, la propuesta es razonable porque ataca los principales riesgos del caso: "
        "liquidez estacional, dependencia de un solo mercado y volatilidad cambiaria. "
        "Tambien mejora la toma de decisiones con indicadores economicos y pruebas de elasticidad."
    )
    doc.add_paragraph(
        "Sin embargo, no elimina por completo la incertidumbre. Eventos externos (aranceles, "
        "cambios politicos, choques globales) pueden alterar los resultados. Por eso, el plan "
        "debe revisarse al menos cada trimestre y ajustarse con datos reales."
    )

    add_heading(doc, "9) Anexos sugeridos (Apartado E)", level=1)
    add_bullet(doc, "INEGI: inflacion, empleo, consumo.")
    add_bullet(doc, "Banxico: tipo de cambio FIX, tasa de referencia, Cetes.")
    add_bullet(doc, "FMI: perspectivas de crecimiento y riesgos globales.")
    add_bullet(doc, "FRED (si aplica): tasas de EE. UU. y variables de apoyo.")
    add_bullet(doc, "Tabla mensual de flujo de caja proyectado vs real.")

    out = r"C:\Users\juanp\Downloads\Propuesta_Caso_Practico_Temporadas_del_Sur.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
