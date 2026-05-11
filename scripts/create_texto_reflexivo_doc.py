from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_center(doc: Document, text: str, size: int = 11, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraphs(doc: Document, paragraphs: list[str]) -> None:
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)


def section_resumen_ejecutivo(doc: Document) -> None:
    add_heading(doc, "1. Resumen Ejecutivo", level=1)
    paragraphs = [
        (
            "A lo largo del curso de Finanzas Personales comprendí que administrar el dinero no es una tarea "
            "aislada ni un conjunto de fórmulas técnicas, sino una práctica cotidiana que impacta directamente "
            "la estabilidad emocional, la libertad de decisión y la posibilidad de construir proyectos de vida "
            "sostenibles. Al inicio del semestre veía las finanzas personales principalmente como una necesidad "
            "operativa: registrar gastos, tratar de ahorrar cuando sobrara dinero y evitar deudas evidentes. "
            "Al cierre del curso, esa percepción cambió por completo, porque entendí que cada decisión financiera "
            "está conectada con variables económicas, riesgos futuros y metas de largo plazo."
        ),
        (
            "El primer concepto que considero esencial es la inflación. Antes la veía como un dato de noticias; "
            "ahora la entiendo como una fuerza real que reduce el poder adquisitivo y que obliga a tomar decisiones "
            "más estratégicas sobre ahorro e inversión. Si una persona guarda dinero sin rendimiento, en términos "
            "reales pierde capacidad de compra con el paso del tiempo. Este aprendizaje transforma la lógica del "
            "ahorro: ya no basta con conservar dinero, es necesario proteger su valor."
        ),
        (
            "El segundo concepto clave es la tasa de interés y su efecto doble sobre la vida financiera. Por una "
            "parte, determina el costo real de créditos y préstamos; por otra, influye en los rendimientos de "
            "instrumentos de ahorro e inversión. Gracias al curso aprendí a analizar no solo la mensualidad de un "
            "crédito, sino también el CAT, las comisiones y la viabilidad de pago dentro de un presupuesto realista. "
            "Esta habilidad evita decisiones impulsivas que, aunque parecen útiles en el corto plazo, pueden generar "
            "presión financiera prolongada."
        ),
        (
            "El tercer concepto significativo es el presupuesto con propósito. Descubrí que presupuestar no implica "
            "restricción extrema, sino claridad y dirección. Un presupuesto bien diseñado permite identificar fugas, "
            "priorizar gastos, asignar recursos a metas concretas y reducir compras impulsivas. También entendí que "
            "el ahorro efectivo no sucede por accidente: se programa desde el inicio del ingreso y se vincula con "
            "objetivos específicos, como un fondo de emergencia, formación profesional o inversión."
        ),
        (
            "El cuarto concepto central es la inversión y la diversificación. El curso me ayudó a romper la idea de "
            "que invertir es solo para personas con grandes capitales o tolerancia extrema al riesgo. Comprendí que "
            "invertir es una herramienta para construir patrimonio y que la diversificación permite equilibrar "
            "seguridad, liquidez y crecimiento. En lugar de concentrar todo en un solo instrumento, una estrategia "
            "financiera sólida distribuye el capital de acuerdo con plazos y objetivos."
        ),
        (
            "El quinto concepto esencial es la previsión financiera integral, que incluye retiro, seguros y "
            "planificación patrimonial. Estos temas ampliaron mi perspectiva porque muestran que las finanzas "
            "personales no se limitan al presente: también consisten en anticiparse a riesgos y proteger a la familia "
            "ante escenarios inciertos. Entendí que planear la pensión desde temprano aprovecha el interés compuesto, "
            "que los seguros protegen el patrimonio frente a eventos inesperados y que ordenar legalmente los bienes "
            "evita conflictos futuros."
        ),
        (
            "En síntesis, estos cinco conceptos forman un sistema de decisión: el entorno macroeconómico condiciona "
            "las opciones; el presupuesto organiza el día a día; el ahorro aporta estabilidad; la inversión impulsa "
            "crecimiento; y la previsión protege el futuro. Por eso considero que el aprendizaje más importante del "
            "curso es asumir una administración financiera consciente, orientada a metas y respaldada por análisis."
        ),
    ]
    add_paragraphs(doc, paragraphs)


def section_analisis_personal(doc: Document) -> None:
    add_heading(doc, "2. Análisis Personal", level=1)
    paragraphs = [
        (
            "Mi percepción sobre las finanzas personales cambió de forma profunda desde el inicio del curso hasta su "
            "cierre. Al principio tenía una visión parcial: pensaba que el éxito financiero dependía, sobre todo, del "
            "nivel de ingreso. Con el desarrollo de las actividades entendí que, aunque el ingreso es importante, la "
            "diferencia real está en la calidad de las decisiones. Personas con ingresos similares pueden tener "
            "resultados muy distintos dependiendo de su capacidad para planear, ahorrar, invertir y gestionar riesgos."
        ),
        (
            "También evolucionó mi forma de entender el consumo. Antes justificaba algunos gastos por recompensa "
            "inmediata, sin evaluar su costo de oportunidad. Ahora me resulta más claro que cada gasto representa una "
            "decisión sobre el uso del tiempo y del esfuerzo que produjo ese dinero. Esta reflexión me llevó a revisar "
            "con mayor profundidad hábitos cotidianos que, aunque parecían pequeños, acumulaban un impacto relevante "
            "sobre mi capacidad de ahorro mensual."
        ),
        (
            "Otro cambio significativo fue reconocer que el crédito no es ni bueno ni malo por sí mismo: su efecto "
            "depende del propósito, del costo y de la capacidad de pago. El curso me ayudó a diferenciar entre deuda "
            "estratégica y deuda de consumo impulsivo. Aprendí a evaluar condiciones como tasa anual, CAT, plazo y "
            "comisiones, y a medir si el compromiso financiero es compatible con un flujo de efectivo saludable."
        ),
        (
            "A partir de estos aprendizajes, identifiqué dos hábitos financieros concretos que planeo implementar "
            "de forma sostenida. El primero es automatizar el ahorro desde el momento en que reciba ingresos. En lugar "
            "de ahorrar al final del mes, apartaré un porcentaje fijo de manera inmediata para construir un fondo de "
            "emergencia y metas de mediano plazo. La lógica es simple: cuando el ahorro se deja al final, casi siempre "
            "se sacrifica; cuando se programa al inicio, se vuelve parte natural de la estructura financiera."
        ),
        (
            "Para consolidar ese hábito, trabajaré con metas específicas y medibles: fondo de emergencia equivalente "
            "a tres meses de gastos básicos en una primera etapa, y posteriormente seis meses como objetivo de mayor "
            "seguridad. Además, separaré cuentas o apartados para evitar mezclar ahorro con gasto operativo. Esta "
            "organización me permitirá visualizar progreso y mantener disciplina, incluso en meses de mayor presión."
        ),
        (
            "El segundo hábito será aplicar un protocolo de decisión antes de contratar cualquier crédito o realizar "
            "compras de monto relevante. Dicho protocolo incluirá tres preguntas: si la compra es necesaria, si puede "
            "pagarse sin comprometer metas prioritarias y si existe una alternativa de menor costo total. En caso de "
            "financiamiento, revisaré el CAT, la tasa efectiva, las comisiones y el impacto en el presupuesto mensual "
            "antes de aceptar cualquier oferta."
        ),
        (
            "Este segundo hábito busca reducir decisiones impulsivas y fortalecer una mentalidad de planeación. Más "
            "que evitar por completo el financiamiento, la meta es usarlo con criterio y solo cuando agregue valor "
            "real al proyecto personal o profesional. En conjunto, ambos hábitos representan una transición de una "
            "administración reactiva hacia una gestión financiera preventiva y orientada a resultados de largo plazo."
        ),
    ]
    add_paragraphs(doc, paragraphs)


def section_aplicacion_practica(doc: Document) -> None:
    add_heading(doc, "3. Aplicación Práctica", level=1)
    paragraphs = [
        (
            "Con base en los conocimientos adquiridos, diseñé un plan financiero personal estructurado en metas por "
            "horizonte temporal, presupuesto mensual, estrategia de ahorro e inversión y mecanismos de seguimiento. "
            "El objetivo general es lograr estabilidad en el corto plazo, crecimiento patrimonial en el mediano plazo "
            "y seguridad financiera para etapas futuras."
        ),
        (
            "Las metas de corto plazo (0 a 12 meses) son: construir un fondo de emergencia equivalente a tres meses "
            "de gastos esenciales, eliminar deudas de consumo de alto costo y mantener un ahorro programado mensual "
            "constante. Las metas de mediano plazo (1 a 5 años) incluyen financiar formación profesional continua, "
            "adquirir activos productivos y consolidar una base de inversión diversificada. Las metas de largo plazo "
            "(más de 5 años) son fortalecer el patrimonio, realizar aportaciones voluntarias al retiro y mantener "
            "protección financiera integral para disminuir vulnerabilidad ante eventos de alto impacto."
        ),
        (
            "El presupuesto mensual propuesto se organiza en cinco bloques. Primero, 50% para necesidades básicas "
            "(vivienda, alimentación, transporte, servicios y salud). Segundo, 20% para ahorro e inversión. Tercero, "
            "10% para educación y desarrollo profesional. Cuarto, 10% para metas personales y gastos discrecionales "
            "controlados. Quinto, 10% para previsión (seguros, mantenimiento y contingencias). Esta distribución no es "
            "rígida, pero funciona como marco de control para evitar desequilibrios."
        ),
        (
            "La estrategia de ahorro contempla automatizar transferencias al inicio de cada periodo de ingreso y "
            "destinar el ahorro en tres capas: fondo de emergencia, metas de mediano plazo y retiro. La primera capa "
            "prioriza liquidez y seguridad; la segunda permite planear objetivos concretos sin recurrir a deuda; la "
            "tercera busca aprovechar el interés compuesto en horizontes largos. Con esta estructura, el ahorro deja "
            "de ser residual y se convierte en una política personal de construcción patrimonial."
        ),
        (
            "Respecto a inversión, el instrumento que incorporaría de manera prioritaria son los CETES, por su bajo "
            "riesgo relativo, respaldo gubernamental y facilidad operativa para comenzar con montos accesibles. Los "
            "utilizaría para objetivos de corto y mediano plazo, comparando siempre el rendimiento real frente a la "
            "inflación. Este análisis es clave: una tasa nominal atractiva no garantiza crecimiento real del patrimonio "
            "si no supera la pérdida de poder adquisitivo."
        ),
        (
            "Como evolución del plan, integraría gradualmente otros instrumentos para diversificar riesgo y mejorar "
            "potencial de rendimiento: fondos de inversión de deuda y renta variable, FIBRAS para exposición "
            "inmobiliaria y una proporción limitada de activos de mayor volatilidad conforme aumente mi conocimiento "
            "y tolerancia al riesgo. La distribución inicial sugerida sería 60% renta fija de bajo riesgo, 25% fondos "
            "diversificados, 10% FIBRAS y 5% activos de crecimiento alto."
        ),
        (
            "Finalmente, el plan incorpora revisión mensual y ajuste trimestral con indicadores simples: porcentaje de "
            "ahorro efectivo, cumplimiento de presupuesto, avance por meta, rendimiento real de inversiones y nivel de "
            "endeudamiento. Esta metodología permite corregir desviaciones a tiempo y mantener coherencia entre metas "
            "financieras y decisiones cotidianas."
        ),
    ]
    add_paragraphs(doc, paragraphs)


def section_reflexion_critica(doc: Document) -> None:
    add_heading(doc, "4. Reflexión Crítica", level=1)
    paragraphs = [
        (
            "Los temas abordados en el curso tienen una aplicación directa en el contexto económico y social actual, "
            "caracterizado por inflación persistente, variaciones en tasas de interés, digitalización acelerada de "
            "servicios financieros y mayor exposición al endeudamiento de consumo. En este escenario, la gestión de "
            "finanzas personales no es un conocimiento accesorio, sino una competencia básica de ciudadanía económica."
        ),
        (
            "Desde una perspectiva crítica, uno de los retos más visibles es la brecha entre acceso a productos "
            "financieros y acceso a educación financiera. Hoy es relativamente sencillo contratar crédito, abrir cuentas "
            "o invertir desde aplicaciones móviles, pero no siempre existen herramientas suficientes para evaluar "
            "riesgos, costos reales y consecuencias de largo plazo. Esta asimetría favorece decisiones rápidas y poco "
            "informadas, especialmente en poblaciones jóvenes o con menor experiencia."
        ),
        (
            "La inflación evidencia con claridad la importancia de esta formación. Cuando los precios aumentan, quienes "
            "no cuentan con presupuesto, ahorro o instrumentos de protección son más vulnerables y suelen recurrir a "
            "deuda cara para sostener gastos básicos. Por el contrario, quienes gestionan su dinero con planeación "
            "pueden ajustar consumo, priorizar necesidades y preservar parte de su estabilidad, incluso en periodos "
            "económicos adversos."
        ),
        (
            "La gestión de finanzas personales también incide en el bienestar individual al reducir incertidumbre y "
            "estrés financiero. Contar con fondo de emergencia, cobertura de seguros y metas de ahorro claras mejora "
            "la capacidad para enfrentar imprevistos sin comprometer por completo el proyecto de vida. En términos de "
            "salud emocional, esta previsión reduce ansiedad y mejora la toma de decisiones, ya que disminuye la "
            "dependencia de soluciones urgentes y costosas."
        ),
        (
            "En el plano colectivo, una población financieramente informada contribuye a una economía más resiliente. "
            "Cuando las personas comparan productos, entienden condiciones de crédito, evitan sobreendeudamiento y "
            "mantienen hábitos de ahorro e inversión, se fortalece la estabilidad de los hogares y disminuyen efectos "
            "sociales asociados a crisis de liquidez familiar. Además, mejora la cultura de prevención y se promueve "
            "una relación más responsable con el sistema financiero."
        ),
        (
            "Otro punto crítico es la planificación de largo plazo. En muchos casos, temas como pensión, seguros o "
            "orden patrimonial se postergan por considerarse lejanos o incómodos. Sin embargo, el curso demuestra que "
            "ignorar estas áreas incrementa la vulnerabilidad futura y puede trasladar costos a la familia o al entorno "
            "social. Planear con anticipación no elimina riesgos, pero sí reduce su impacto y permite responder con "
            "mayor dignidad y autonomía."
        ),
        (
            "Como conclusión, la gestión de finanzas personales debe entenderse como una práctica ética y estratégica: "
            "ética, porque involucra responsabilidad sobre recursos que sostienen proyectos propios y familiares; "
            "estratégica, porque exige decisiones informadas en un entorno cambiante. Integrar lo aprendido en el curso "
            "implica pasar de la improvisación a la planeación, del consumo reactivo al uso consciente del dinero y de "
            "la incertidumbre permanente a una construcción gradual de bienestar individual y colectivo."
        ),
    ]
    add_paragraphs(doc, paragraphs)


def main() -> None:
    doc = Document()

    add_title(doc, "Actividad Final - Texto Reflexivo")
    add_center(doc, "Curso: Finanzas Personales", size=12, bold=True)
    add_center(doc, "Estudiante: [Tu nombre completo]", size=11)
    add_center(doc, f"Fecha: {date.today().strftime('%d/%m/%Y')}", size=11)
    doc.add_page_break()

    section_resumen_ejecutivo(doc)
    doc.add_page_break()
    section_analisis_personal(doc)
    doc.add_page_break()
    section_aplicacion_practica(doc)
    doc.add_page_break()
    section_reflexion_critica(doc)

    out = (
        r"C:\Users\juanp\Desktop\Materias\10mo semestre\Tareas\Finanzas Personales"
        r"\Actividad Final - Texto Reflexivo Integrado.docx"
    )
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
