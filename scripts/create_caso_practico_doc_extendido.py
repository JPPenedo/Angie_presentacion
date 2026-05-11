from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_center(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def main():
    doc = Document()

    # Portada
    add_title(doc, "Fundamentos de Economia y Finanzas")
    add_center(doc, "Caso practico integral", size=14, bold=True)
    add_center(doc, "Primer reto profesional: Finanzas entre temporadas", size=12)
    doc.add_paragraph("")
    add_center(doc, "Propuesta para: Director General de Temporadas del Sur", size=11)
    add_center(doc, "Elaborado por: [Tu nombre completo]", size=11)
    add_center(doc, "Curso: Fundamentos de Economia y Finanzas", size=11)
    add_center(doc, "Fecha: " + date.today().strftime("%d/%m/%Y"), size=11)
    doc.add_page_break()

    add_heading(doc, "1) Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "Temporadas del Sur es una empresa que funciona bien en temporada alta, pero sufre "
        "durante varios meses del anio por falta de liquidez. El problema de fondo no es solo "
        "vender mas, sino administrar mejor el dinero entre una temporada y otra."
    )
    doc.add_paragraph(
        "En este caso, el riesgo no viene de un solo lado: hay dependencia fuerte de Estados Unidos, "
        "posibles cambios de aranceles, inflacion, variaciones en el tipo de cambio y costos de diesel. "
        "Por eso la propuesta no se centra en una sola accion, sino en un paquete de medidas que se "
        "refuerzan entre si."
    )
    doc.add_paragraph(
        "El enfoque combina tres ideas practicas: ordenar el flujo de caja desde ahora, reducir riesgos "
        "que pegan directo al margen (tipo de cambio y costos), y abrir camino para vender en otros "
        "mercados en el mediano plazo."
    )

    add_heading(doc, "2) Situacion y contexto economico (Apartado A)", level=1)
    doc.add_paragraph(
        "La empresa estima ventas por 3.6 millones de pesos al anio. Si el 70% se va a produccion y "
        "10% a gastos administrativos, queda poco espacio para absorber errores. Esto explica por que "
        "una mala semana de caja o un movimiento fuerte en el dolar puede complicar toda la operacion."
    )
    doc.add_paragraph(
        "Otro punto importante es la concentracion de ventas: 70% ocurre en agosto y septiembre. "
        "Eso significa que una decision tomada tarde (por ejemplo, no asegurar inventario o no cubrir "
        "tipo de cambio) puede afectar justo los meses donde se juega gran parte del resultado anual."
    )
    doc.add_paragraph("Conceptos economicos aplicados al caso, en terminos simples:")
    add_bullet(
        doc,
        "Oferta y demanda: la demanda de gorros y guantes sube con frio y cae en meses calidos. "
        "Tiene sentido planear produccion y compras por temporada, no con un promedio anual plano.",
    )
    add_bullet(
        doc,
        "Elasticidad precio: no todos los clientes reaccionan igual a una subida de precio. "
        "Medir elasticidad evita subir precios a ciegas y perder volumen innecesariamente.",
    )
    add_bullet(
        doc,
        "Inflacion: presiona costos de insumos y transporte. Ajustar precios sin analizar demanda "
        "puede resolver margen en papel, pero bajar ventas reales.",
    )
    add_bullet(
        doc,
        "Tasas de interes: cuando suben, el credito pesa mas. Conviene negociar lineas antes de una "
        "emergencia para no aceptar condiciones caras por urgencia.",
    )
    add_bullet(
        doc,
        "Tipo de cambio: si la empresa vende en dolares pero paga costos en pesos y algunos insumos "
        "ligados a dolar, la volatilidad puede jugar a favor o en contra rapidamente.",
    )
    doc.add_paragraph(
        "En conjunto, el contexto economico actual justifica una estrategia financiera mas activa. "
        "No basta con producir y esperar pedidos: hay que gestionar riesgo y liquidez de forma preventiva."
    )

    add_heading(doc, "3) Propuestas de corto plazo (Apartado C)", level=1)
    add_heading(doc, "3.1 Flujo de caja y liquidez inmediata", level=2)
    add_bullet(
        doc,
        "Implementar presupuesto semanal de caja por 6 meses. Esto tiene sentido porque la empresa "
        "opera con estacionalidad fuerte; un seguimiento mensual llega tarde para corregir.",
    )
    add_bullet(
        doc,
        "Negociar con proveedores ampliar plazo de 45 a 60 dias en meses de baja demanda. "
        "Tiene sentido porque libera efectivo justo cuando menos entra dinero por ventas.",
    )
    add_bullet(
        doc,
        "Abrir una linea revolvente para capital de trabajo y usarla solo cuando el flujo lo requiera. "
        "Tiene sentido porque da flexibilidad sin convertir el credito en gasto permanente.",
    )
    add_bullet(
        doc,
        "Reducir operaciones en efectivo y mover pagos a banca digital. "
        "Tiene sentido por seguridad, trazabilidad y control contable.",
    )
    doc.add_paragraph(
        "Estas acciones son realistas porque no dependen de cambiar todo el modelo de negocio. "
        "Se pueden iniciar de inmediato y mejoran la capacidad de reaccion en semanas, no en anios."
    )

    add_heading(doc, "3.2 Cobros y pagos: orden operativo", level=2)
    add_bullet(
        doc,
        "Solicitar anticipo del 30% en pedidos grandes de exportacion. "
        "Tiene sentido porque traslada parte del riesgo al inicio del proceso y reduce presion de capital.",
    )
    add_bullet(
        doc,
        "Ofrecer descuento pequeno por pronto pago solo a clientes de bajo riesgo. "
        "Tiene sentido porque acelera entradas de caja sin sacrificar demasiado margen.",
    )
    add_bullet(
        doc,
        "Priorizar pagos criticos (nomina, proveedores clave) y calendarizar no criticos. "
        'Tiene sentido porque evita cortar la operacion por falta de pago en rubros esenciales.',
    )
    doc.add_paragraph(
        "El objetivo aqui no es solo 'cobrar mas rapido', sino que el dinero entre y salga de forma "
        "coordinada para evitar picos de tension financiera."
    )

    add_heading(doc, "3.3 Cobertura cambiaria y costos de energia", level=2)
    add_bullet(
        doc,
        "Cubrir entre 50% y 70% de los flujos esperados en dolares a 3-6 meses con forwards. "
        "Tiene sentido porque protege margen sin bloquear totalmente posibles beneficios del mercado.",
    )
    add_bullet(
        doc,
        "Definir una regla interna para activar cobertura adicional si el tipo de cambio supera umbral "
        "predefinido. Tiene sentido porque convierte una decision emocional en una decision tecnica.",
    )
    add_bullet(
        doc,
        "Analizar cobertura parcial de diesel si el gasto representa 10% de costos. "
        "Tiene sentido porque una subida sostenida de energia pega directo al costo unitario.",
    )
    doc.add_paragraph(
        "Cubrir riesgo no es especular. Es fijar piso de certidumbre para proteger utilidad y cumplir "
        "compromisos de produccion."
    )

    add_heading(doc, "4) Propuestas de largo plazo (Apartado B)", level=1)
    add_heading(doc, "4.1 Sistema de monitoreo economico", level=2)
    add_bullet(
        doc,
        "Crear tablero mensual con inflacion, tipo de cambio, tasas, empleo y consumo. "
        "Tiene sentido porque estas variables anticipan cambios en costos y demanda.",
    )
    add_bullet(
        doc,
        "Asignar responsable y fecha fija de revision mensual. "
        "Tiene sentido porque un indicador sin seguimiento termina sin uso practico.",
    )
    add_bullet(
        doc,
        "Definir semaforos (verde, amarillo, rojo) con acciones concretas. "
        "Tiene sentido porque permite pasar de diagnostico a ejecucion.",
    )

    add_heading(doc, "4.2 Diversificacion de mercados", level=2)
    doc.add_paragraph(
        "La dependencia de EE. UU. (85% de ventas) es rentable mientras todo va bien, pero es riesgosa "
        "cuando hay cambios regulatorios o politicos. Diversificar no significa abandonar ese mercado, "
        "sino reducir la vulnerabilidad."
    )
    add_bullet(
        doc,
        "Iniciar con Colombia y Espana por cercania cultural/comercial. "
        "Tiene sentido porque reduce friccion de entrada y permite aprender mas rapido.",
    )
    add_bullet(
        doc,
        "Entrar a Brasil y Alemania en una segunda etapa con pilotos acotados. "
        "Tiene sentido porque son mercados atractivos, pero mas exigentes en logistica y regulacion.",
    )
    add_bullet(
        doc,
        "Meta a 3 anios: bajar participacion de EE. UU. de 85% a 65%. "
        "Tiene sentido porque mejora balance de riesgo sin romper estructura comercial actual.",
    )
    doc.add_paragraph(
        "Diversificar tambien ayuda a estabilizar ventas anuales: si un mercado se frena, otro puede "
        "compensar parcialmente."
    )

    add_heading(doc, "4.3 Ahorro e inversion de excedentes", level=2)
    add_bullet(
        doc,
        "Colocar excedentes temporales en instrumentos de alta liquidez y bajo riesgo (Cetes/fondos deuda CP). "
        "Tiene sentido porque evita que el efectivo pierda valor por inflacion.",
    )
    add_bullet(
        doc,
        "Construir fondo de contingencia de al menos 2 meses de gastos fijos. "
        "Tiene sentido porque reduce dependencia de deuda cara en momentos de estres.",
    )

    add_heading(doc, "5) Analisis de mercado, equilibrio y elasticidad", level=1)
    doc.add_paragraph(
        "Para fijar precios con mayor precision, la empresa necesita medir como responde la demanda. "
        "La forma mas util es experimentar de manera controlada y aprender con datos propios."
    )
    add_bullet(
        doc,
        "Aplicar pruebas de precio (+/-5%) por canal y ciudad durante 4-6 semanas. "
        "Tiene sentido porque permite ver reaccion real sin comprometer toda la cartera.",
    )
    add_bullet(
        doc,
        "Comparar unidades vendidas, ticket promedio y margen por prueba. "
        "Tiene sentido porque vender mas no siempre significa ganar mas.",
    )
    add_bullet(
        doc,
        "Calcular elasticidad-precio de la demanda por mercado. "
        "Tiene sentido porque un mismo precio no funciona igual en todos los paises.",
    )
    add_bullet(
        doc,
        "Usar resultados para aproximar precio de equilibrio por segmento. "
        "Tiene sentido porque evita decisiones uniformes en mercados distintos.",
    )
    doc.add_paragraph(
        "Este punto es clave para defender cualquier ajuste de precio frente a la direccion: "
        "se decide con evidencia y no solo con intuicion."
    )

    add_heading(doc, "6) Simulacion financiera orientativa", level=1)
    doc.add_paragraph("Escenario base anual del caso:")
    add_bullet(doc, "Ventas: 3.6 millones MXN")
    add_bullet(doc, "Costos de produccion (70%): 2.52 millones MXN")
    add_bullet(doc, "Gastos administrativos (10%): 0.36 millones MXN")
    add_bullet(doc, "Margen antes de gastos financieros e impuestos: 0.72 millones MXN")
    doc.add_paragraph(
        "Este margen es sensible. Un aumento de costos o una baja de demanda puede comerse una parte "
        "importante de la utilidad esperada."
    )
    doc.add_paragraph("Escenario estresado razonable:")
    add_bullet(doc, "Inflacion mayor a lo esperado y alza en costos de transporte/insumos.")
    add_bullet(doc, "Atraso de cobros o pedidos menores en exportacion.")
    add_bullet(doc, "Necesidad de credito de corto plazo para sostener operacion.")
    doc.add_paragraph(
        "Por eso tiene sentido combinar cobertura, mejor cobranza y reserva de liquidez. "
        "Ninguna medida sola resuelve todo."
    )

    add_heading(doc, "7) Seguimiento y evaluacion de resultados", level=1)
    doc.add_paragraph(
        "Para saber si la estrategia funciona, se recomienda un corte mensual y uno trimestral "
        "con indicadores simples."
    )
    add_bullet(
        doc,
        "Ventas reales vs presupuestadas (por mercado y por canal). "
        "Tiene sentido porque muestra donde se desvia el plan.",
    )
    add_bullet(
        doc,
        "Margen bruto real vs objetivo. "
        "Tiene sentido porque permite ver si el crecimiento realmente deja rentabilidad.",
    )
    add_bullet(
        doc,
        "Flujo de caja neto semanal. "
        "Tiene sentido porque anticipa faltantes antes de que se vuelvan problema.",
    )
    add_bullet(
        doc,
        "Porcentaje de exposicion cambiaria cubierta. "
        "Tiene sentido porque mide disciplina de gestion de riesgo.",
    )
    add_bullet(
        doc,
        "Peso de ventas fuera de EE. UU. "
        "Tiene sentido porque evalua avance real de diversificacion.",
    )
    doc.add_paragraph(
        "Como metrica de error para proyecciones se puede usar MAPE (Error Porcentual Medio Absoluto). "
        "Es facil de interpretar y util para ir ajustando supuestos."
    )

    add_heading(doc, "8) Conclusiones (respuesta directa al caso)", level=1)
    doc.add_paragraph(
        "Si, la propuesta es razonable para el caso porque responde de forma concreta a los tres "
        "riesgos principales: estacionalidad de caja, concentracion de mercado y volatilidad de costos/"
        "tipo de cambio."
    )
    doc.add_paragraph(
        "Tambien es una propuesta viable porque prioriza acciones que se pueden empezar ya (orden de caja, "
        "politica de cobros y coberturas parciales), y al mismo tiempo construye capacidades para el futuro "
        "(monitoreo economico y diversificacion)."
    )
    doc.add_paragraph(
        "No es una solucion perfecta, porque siempre habra eventos externos que la empresa no controla "
        "(aranceles, decisiones de politica economica, choques globales). Aun asi, da una base mucho mas "
        "solida para decidir con menos improvisacion y mayor control financiero."
    )

    add_heading(doc, "9) Plan de implementacion sugerido", level=1)
    doc.add_paragraph("Primeros 90 dias:")
    add_bullet(doc, "Semana 1-2: presupuesto de caja semanal y orden de pagos.")
    add_bullet(doc, "Semana 3-4: negociacion con proveedores y bancos.")
    add_bullet(doc, "Mes 2: politica de cobertura cambiaria y tablero economico.")
    add_bullet(doc, "Mes 3: prueba piloto de elasticidad en al menos dos canales.")
    doc.add_paragraph("Horizonte 6 a 12 meses:")
    add_bullet(doc, "Abrir primer mercado alternativo y medir desempeno comercial.")
    add_bullet(doc, "Consolidar fondo de contingencia.")
    add_bullet(doc, "Ajustar precios y cobertura con resultados reales.")

    add_heading(doc, "10) Anexos sugeridos (Apartado E)", level=1)
    add_bullet(doc, "INEGI: inflacion, empleo y consumo privado.")
    add_bullet(doc, "Banxico: tipo de cambio FIX, tasa de referencia, Cetes.")
    add_bullet(doc, "FMI: perspectivas de crecimiento para Mexico, EE. UU. y paises objetivo.")
    add_bullet(doc, "Series historicas propias de ventas, margen y rotacion de inventario.")
    add_bullet(doc, "Formato de seguimiento mensual (presupuesto vs real).")

    out = r"C:\Users\juanp\Downloads\Propuesta_Caso_Practico_Temporadas_del_Sur_Extendida.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
