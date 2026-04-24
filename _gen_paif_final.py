"""
Generador del documento PAIF Final (Equipo 3 · ODS 16 · Denuncia Verde).

Objetivo: producir un .docx que cumpla con TODOS los criterios de la rúbrica
`Rúbrica PAIF final 202610`, respetando el formato de portada que ha usado
el equipo en actividades anteriores (Anáhuac · Responsabilidad Social y
Sustentabilidad).

Uso:
    .\venv\Scripts\python.exe _gen_paif_final.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- Rutas ----------
DEST_DIR = Path(
    r"C:\Users\juanp\Desktop\Materias\10mo semestre\Tareas\Responsabilidad social y sutentabilidad"
)
FILENAME = "PAIF Final. Equipo 3 ODS16 - Denuncia Verde.docx"


# ---------- Datos ----------
EQUIPO = [
    {"nombre": "Juan Pablo Penedo Antúnez",       "correo": "juan.penedo@anahuac.mx",      "carrera": "Actuaría"},
    {"nombre": "David Manuel Ruiz Pérez",         "correo": "david.ruiz11@anahuac.mx",     "carrera": "Ingeniería Industrial"},
    {"nombre": "Leonardo Mayoral García",         "correo": "leonardo.mayoral@anahuac.mx", "carrera": "Administración y Finanzas"},
    {"nombre": "Pablo Andrés Pacheco Colmenares", "correo": "pablo.pacheco@anahuac.mx",    "carrera": "Derecho"},
]
DOCENTE = "Prof. Clemente Sánchez Uribe"
FECHA_ENTREGA = "24/04/26"
PROYECTO = "Denuncia Verde · Plataforma comunitaria contra la corrupción ambiental"


# ---------- Helpers de formato ----------
def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None, underline=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    # Ensure asiatic font uses same family (Word quirk)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def add_para(doc, text="", size=11, bold=False, italic=False, align=None, space_after=6, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, name="Calibri", size=16, bold=True, color=(31, 73, 42))
    # underline-like border bottom
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "7DA27C")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, name="Calibri", size=13, bold=True, color=(47, 107, 59))
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, name="Calibri", size=12, bold=True, italic=True, color=(57, 85, 66))
    return p


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        run = p.add_run(bold_lead + ". ")
        set_font(run, size=11, bold=True)
        run2 = p.add_run(text)
        set_font(run2, size=11)
    else:
        run = p.add_run(text)
        set_font(run, size=11)
    return p


def add_justified(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0.75)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(doc, headers, rows, widths=None, header_color="2F6B3B"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, size=10, bold=True, color=(255, 255, 255))
        shade_cell(cell, header_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Rows
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_font(run, size=10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return table


def add_page_break(doc):
    doc.add_page_break()


# ---------- Construcción del documento ----------
doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Estilo base
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


# ---------- PORTADA (6 elementos) ----------
add_para(doc, "UNIVERSIDAD ANÁHUAC", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, "", space_after=60)

add_para(doc, "MATERIA EN LÍNEA", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "RESPONSABILIDAD SOCIAL Y SUSTENTABILIDAD", size=13, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

add_para(doc, "Proyecto de Aplicación Integradora Final (PAIF)", size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para(doc, "Equipo 3 — Paz, justicia e instituciones sólidas (ODS 16)", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para(doc, f"Proyecto: {PROYECTO}", size=12, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=28)

for m in EQUIPO:
    add_para(doc, f"{m['nombre']}  —  {m['correo']}", size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

add_para(doc, "", space_after=18)
add_para(doc, f"DOCENTE: {DOCENTE}", size=11, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

add_para(doc, "Fecha de entrega", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, FECHA_ENTREGA, size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

add_page_break(doc)


# ---------- ÍNDICE simplificado ----------
add_h1(doc, "Índice")
indice = [
    "Introducción",
    "1. Análisis del problema y entorno",
    "2. Fundamentación teórica",
    "3. Objetivo general y objetivos específicos",
    "4. Alineación con la ISO 26000",
    "5. Doctrina Social de la Iglesia, ética y transparencia",
    "6. Uso de recursos y sostenibilidad",
    "7. Derechos humanos",
    "8. Matriz de grupos de interés y reflexión sobre la transformación social",
    "9. Participación y colaboración del equipo",
    "10. Referencias",
]
for line in indice:
    add_para(doc, line, size=11, space_after=2)

add_page_break(doc)


# ---------- INTRODUCCIÓN ----------
add_h1(doc, "Introducción")
add_justified(doc,
    "Este documento integra el Proyecto de Aplicación Integradora Final (PAIF) del Equipo 3 "
    "en la materia Responsabilidad Social y Sustentabilidad. El proyecto, titulado “Denuncia Verde”, "
    "es una plataforma comunitaria que permite reportar de forma segura y anónima la corrupción "
    "ambiental —tala ilegal, uso indebido del agua, contaminación y otras prácticas que deterioran "
    "los ecosistemas locales— y, al mismo tiempo, ofrece contenidos de educación socioambiental y "
    "difusión de derechos. Se enmarca en el ODS 16 (Paz, justicia e instituciones sólidas), con aporte "
    "directo a los ODS 13 (Acción por el clima) y 15 (Vida de ecosistemas terrestres).")
add_justified(doc,
    "La intervención se sitúa en Querétaro, México, territorio en el que la presión sobre bosques, "
    "acuíferos y suelos de conservación convive con irregularidades documentadas por PROFEPA y "
    "organizaciones civiles. Nuestro objetivo no es sustituir a la autoridad, sino aumentar la "
    "rendición de cuentas por la vía ciudadana, articulándonos con instituciones públicas, "
    "universidades y organizaciones ambientales.")
add_justified(doc,
    "A lo largo del documento se presentan: el análisis del problema y entorno; la fundamentación "
    "teórica con aportes de Klitgaard, Sen, Carroll, Porter & Kramer y Salamon; los objetivos; la "
    "alineación con la ISO 26000 (siete materias y siete principios); dos principios de la Doctrina "
    "Social de la Iglesia; la planeación de recursos con su contribución sostenible; los derechos "
    "humanos atendidos; la matriz de grupos de interés con análisis desde las profesiones del equipo; "
    "y finalmente la reflexión individual de los integrantes.")

add_page_break(doc)


# ---------- 1. ANÁLISIS DEL PROBLEMA ----------
add_h1(doc, "1. Análisis del problema y entorno")

add_h2(doc, "1.1 Contexto y alineación con el eje temático")
add_justified(doc,
    "El eje temático asignado al equipo es el ODS 16: Paz, justicia e instituciones sólidas, con "
    "énfasis en la meta 16.5 (reducir sustancialmente la corrupción y el soborno). Seleccionamos "
    "como problema específico la corrupción ambiental en México, particularmente en el estado de "
    "Querétaro, entendida como el uso ilegítimo del poder público o privado para obtener beneficios "
    "a costa de los ecosistemas y los bienes comunes (agua, suelo, fauna, bosques).")
add_justified(doc,
    "Esta problemática es relevante para la comunidad: Querétaro es uno de los estados de mayor "
    "crecimiento económico del país, lo que ha puesto en tensión al acuífero del Valle, a zonas de "
    "conservación como la Peña de Bernal y El Zamorano, y a la Reserva de la Sierra Gorda. La "
    "PROFEPA ha documentado más de 3,000 irregularidades ambientales en la zona centro del país "
    "entre 2019 y 2024, y Transparencia Mexicana registra a Querétaro entre los estados con "
    "denuncias recurrentes por permisos ambientales cuestionados.")

add_h2(doc, "1.2 El problema desde tres dimensiones")
add_h3(doc, "Dimensión social")
add_justified(doc,
    "La impunidad en los temas ambientales erosiona la confianza ciudadana en las instituciones y "
    "debilita el tejido social de comunidades vecinas a zonas afectadas. Muchas personas identifican "
    "irregularidades pero no denuncian por miedo a represalias, por desconocimiento del procedimiento "
    "o por la percepción —muchas veces acertada— de que el reporte no tendrá efecto.")
add_h3(doc, "Dimensión económica")
add_justified(doc,
    "La corrupción ambiental genera costos públicos —remediación de suelos y cuerpos de agua, "
    "pérdidas agrícolas, atención a salud pública— que terminan pagando los propios ciudadanos. A "
    "la vez, distorsiona la competencia: las empresas que cumplen la regulación quedan en desventaja "
    "frente a quienes evaden permisos o tolerancia ilegal.")
add_h3(doc, "Dimensión ambiental")
add_justified(doc,
    "La evidencia es directa: tala ilegal, sobreexplotación de acuíferos, cambios de uso de suelo "
    "sin consulta, vertimiento de residuos en arroyos. La corrupción acelera la pérdida de "
    "biodiversidad y deteriora servicios ecosistémicos de los que depende la ciudad.")

add_h2(doc, "1.3 Viabilidad y evidencias")
add_justified(doc,
    "El proyecto es viable técnicamente: la plataforma web funciona con tecnologías estándar "
    "(Django, HTML/CSS/JS) y almacenamiento local en el navegador para la versión piloto, lo que "
    "reduce costos y permite desplegarla con un presupuesto inicial de $40,000–$45,000 MXN. Es viable "
    "socialmente: Querétaro cuenta con universidades, ONGs como Colectivo Pro Bosques y una base "
    "ciudadana sensibilizada tras conflictos públicos por temas de agua y suelo. Finalmente, es "
    "viable institucionalmente: la PROFEPA, la CONAGUA y la SEDESU estatal ya operan canales oficiales "
    "con los que la plataforma puede articular denuncias canalizadas. La necesidad está documentada "
    "(PROFEPA, 2024; Transparencia Mexicana, 2023) y el impacto positivo potencial —más vigilancia "
    "ciudadana, menos impunidad, mayor cultura ambiental— es coherente con la meta ODS 16.5.")

add_page_break(doc)


# ---------- 2. FUNDAMENTACIÓN TEÓRICA ----------
add_h1(doc, "2. Fundamentación teórica")

add_h2(doc, "2.1 Corrupción e institucionalidad")
add_justified(doc,
    "Robert Klitgaard (1988, 1998) formaliza la corrupción como el resultado del monopolio de "
    "decisión más la discrecionalidad, menos la rendición de cuentas (C = M + D − A). El modelo es "
    "iluminador para nuestro proyecto: no pretendemos modificar el monopolio (M) ni eliminar la "
    "discrecionalidad (D), que son propias de toda autoridad regulatoria; sí buscamos aumentar A "
    "mediante vigilancia ciudadana sistemática. Cada denuncia registrada en la plataforma es un "
    "incremento marginal en A.")

add_h2(doc, "2.2 Desarrollo como libertad")
add_justified(doc,
    "Amartya Sen (1999) propone entender el desarrollo como la expansión de las libertades reales "
    "de las personas, y a la corrupción como una privación de libertad. Desde esa mirada, cuando "
    "un funcionario tolera la tala ilegal en un bosque de uso común está restringiendo la libertad "
    "de la comunidad para disfrutar de ese bien y tomar decisiones sobre él. La plataforma amplía "
    "libertades al dotar a la ciudadanía de un canal de expresión y vigilancia.")

add_h2(doc, "2.3 Responsabilidad Social como pirámide")
add_justified(doc,
    "Archie Carroll (1991) estructura la responsabilidad social en cuatro niveles: económico, "
    "legal, ético y filantrópico. Nuestro proyecto se ubica en los niveles legal y ético: no "
    "hacemos filantropía (donativos), sino que contribuimos a que el sistema legal ambiental "
    "funcione mejor, y lo hacemos desde un marco ético que privilegia la protección del "
    "denunciante y la verificación responsable.")

add_h2(doc, "2.4 Valor compartido y aliados")
add_justified(doc,
    "Michael Porter y Mark Kramer (2011) proponen la creación de valor compartido: políticas y "
    "prácticas que mejoran la competitividad de una organización mientras simultáneamente benefician "
    "a la sociedad. Aplicado a Denuncia Verde, el valor compartido se expresa en la alianza con "
    "empresas y gobierno: ambos ganan con reglas claras, competencia leal y recursos naturales "
    "disponibles a largo plazo.")

add_h2(doc, "2.5 Sociedad civil organizada")
add_justified(doc,
    "Lester Salamon (2004) documenta el rol creciente del tercer sector en la provisión de bienes "
    "públicos y en la vigilancia democrática, especialmente cuando el Estado tiene capacidades "
    "limitadas. La plataforma se entiende como un ejercicio de sociedad civil organizada que "
    "complementa —no sustituye— a la autoridad ambiental.")

add_h2(doc, "2.6 Utilidad y beneficios del proyecto")
add_bullet(doc, "Canal anónimo que reduce el costo personal de denunciar y aumenta la tasa de reportes útiles.", bold_lead="Para la ciudadanía")
add_bullet(doc, "Información sistematizada sobre patrones de irregularidad (tipo, lugar, frecuencia) para priorizar inspecciones.", bold_lead="Para la autoridad")
add_bullet(doc, "Evidencia pública que reduce la impunidad estructural y eleva el costo político de tolerar irregularidades.", bold_lead="Para el entorno")
add_bullet(doc, "Un modelo replicable: se puede escalar a otros estados o a otros ejes (violencia, servicios).", bold_lead="Para la academia")

add_page_break(doc)


# ---------- 3. OBJETIVOS ----------
add_h1(doc, "3. Objetivo general y objetivos específicos")

add_h2(doc, "3.1 Pregunta de transformación social")
add_justified(doc,
    "¿Cómo fortalecer, desde la ciudadanía, los mecanismos de transparencia y rendición de cuentas "
    "en la gestión de los recursos naturales de Querétaro, para reducir la corrupción ambiental y "
    "proteger ecosistemas y comunidades locales?")

add_h2(doc, "3.2 Objetivo general")
add_justified(doc,
    "Diseñar e implementar una plataforma digital comunitaria —Denuncia Verde— que permita "
    "reportar de forma segura, anónima y verificable las prácticas de corrupción que afectan "
    "al medio ambiente en Querétaro, acompañada de contenidos educativos y de difusión de "
    "derechos, con el fin de incrementar la rendición de cuentas y contribuir de manera concreta "
    "al cumplimiento de la meta 16.5 de los Objetivos de Desarrollo Sostenible.")

add_h2(doc, "3.3 Objetivos específicos")
add_bullet(doc,
    "Desarrollar un canal digital anónimo y público de denuncia ambiental con al menos 6 categorías "
    "(tala ilegal, uso indebido del agua, uso indebido del suelo, contaminación, daño a fauna o "
    "flora, otra) y articulado con PROFEPA, CONAGUA, SEDESU Querétaro y la CEDH estatal.",
    bold_lead="Social")
add_bullet(doc,
    "Publicar seis micro-módulos educativos abiertos sobre corrupción ambiental, derechos humanos, "
    "cómo denunciar, recursos naturales locales, participación ciudadana y responsabilidad social "
    "(ISO 26000), con lenguaje accesible y fuentes verificables.",
    bold_lead="Ambiental")
add_bullet(doc,
    "Operar el proyecto con un presupuesto piloto máximo de $45,000 MXN en seis meses, "
    "privilegiando tecnologías de bajo costo y voluntariado universitario, para demostrar que el "
    "modelo es replicable y económicamente sostenible en otras regiones.",
    bold_lead="Económico")

add_page_break(doc)


# ---------- 4. ISO 26000 ----------
add_h1(doc, "4. Alineación con los principios y materias fundamentales de la ISO 26000")

add_h2(doc, "4.1 Responsabilidad Social, no filantropía")
add_justified(doc,
    "Denuncia Verde no entrega donativos ni asistencia puntual. Actúa sobre las causas "
    "estructurales de un problema social (la corrupción ambiental) a través de un cambio en las "
    "reglas de juego: más transparencia, más información ciudadana, más articulación institucional. "
    "Esto la ubica inequívocamente dentro de la Responsabilidad Social en el sentido de la ISO 26000 "
    "y no dentro de la filantropía tradicional.")

add_h2(doc, "4.2 Las siete materias fundamentales")
materias = [
    ("1. Gobernanza organizacional",
     "El proyecto define roles, responsabilidades y procesos claros: coordinador general, "
     "asesor jurídico ambiental, desarrollador, especialista educativo. Cada denuncia se trata "
     "según un protocolo documentado (recepción → validación → publicación anónima → canalización "
     "opcional).",
     True),
    ("2. Derechos humanos",
     "Se atienden tres derechos fundamentales: a un medio ambiente sano (ONU, 2022), a la "
     "participación ciudadana (DUDH, art. 21) y al acceso a la información (DUDH, art. 19). La "
     "protección del denunciante es un principio rector de la plataforma.",
     True),
    ("3. Prácticas laborales",
     "Aunque el piloto opera con voluntariado universitario, se establecen límites horarios y "
     "reconocimiento formal de las horas dedicadas (servicio social); en la fase de escalamiento "
     "se buscará formalizar contrataciones dignas.",
     False),
    ("4. Medio ambiente",
     "Es la materia central del proyecto: la plataforma existe para proteger recursos naturales "
     "de la corrupción que los erosiona. Las denuncias se categorizan por tipo de impacto ambiental.",
     True),
    ("5. Prácticas justas de operación",
     "Los convenios con autoridades, universidades y ONGs se rigen por principios de transparencia "
     "y reciprocidad; no hay remuneraciones encubiertas ni relaciones de exclusividad abusivas.",
     True),
    ("6. Asuntos de consumidores",
     "Aplicado analógicamente a las personas usuarias de la plataforma: se les informa con "
     "claridad sobre el tratamiento de sus datos, se les garantiza anonimato por defecto y se les "
     "ofrece educación cívica sin publicidad ni manipulación.",
     False),
    ("7. Participación activa y desarrollo de la comunidad",
     "El proyecto es, en esencia, un dispositivo de participación comunitaria: los ciudadanos no "
     "son receptores pasivos de información, sino vigilantes activos de su entorno y actores en "
     "los talleres formativos.",
     True),
]
add_justified(doc,
    "A continuación se explican las siete materias y se marca con ✓ las cuatro que están integradas "
    "de forma explícita y estructural en el proyecto:")
for titulo, desc, activa in materias:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run1 = p.add_run(("✓ " if activa else "· ") + titulo + ". ")
    set_font(run1, size=11, bold=True, color=(47, 107, 59) if activa else (90, 90, 90))
    run2 = p.add_run(desc)
    set_font(run2, size=11)

add_h2(doc, "4.3 Los siete principios y su aplicación")
principios = [
    ("Rendición de cuentas",
     "El panel público de denuncias es visible para cualquier persona; el equipo publica "
     "trimestralmente un reporte con el número de denuncias recibidas, su categoría y su estatus "
     "(recibida, en seguimiento, canalizada a autoridad).",
     True),
    ("Transparencia",
     "Los módulos educativos declaran explícitamente sus fuentes (ONU, SEMARNAT, PROFEPA, autores "
     "académicos citados). El código fuente del prototipo puede ser auditado por la comunidad "
     "universitaria.",
     True),
    ("Comportamiento ético",
     "Anonimato por defecto, prohibición de explotación comercial de los datos, compromiso de no "
     "difamar ni acusar individualmente sin pruebas. El equipo firma una declaración de principios "
     "éticos basada en la DSI y la ISO 26000.",
     True),
    ("Respeto a los intereses de las partes interesadas",
     "La matriz de grupos de interés (capítulo 8) formaliza este principio: cada actor tiene un "
     "canal, un rol y un impacto reconocido.",
     False),
    ("Respeto al principio de legalidad",
     "Toda denuncia canalizada se envía conforme a los procedimientos oficiales (PROFEPA, CONAGUA, "
     "SEDESU). La plataforma no suplanta funciones ni invade competencias de autoridad.",
     True),
    ("Respeto a la normativa internacional",
     "Alineación explícita con los ODS 13, 15 y 16 de Naciones Unidas y con el Acuerdo de Escazú "
     "(2018), que garantiza en América Latina el acceso a la información, la participación y la "
     "justicia ambiental.",
     False),
    ("Respeto a los derechos humanos",
     "El derecho a un medio ambiente sano es el eje del proyecto; el derecho a la participación y a "
     "la información son el medio. Ningún usuario puede ser forzado a revelar su identidad.",
     False),
]
add_justified(doc,
    "Se marca con ✓ los cuatro principios cuya aplicación es más directa y estructural en el "
    "proyecto:")
for nombre, desc, activa in principios:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run1 = p.add_run(("✓ " if activa else "· ") + nombre + ". ")
    set_font(run1, size=11, bold=True, color=(47, 107, 59) if activa else (90, 90, 90))
    run2 = p.add_run(desc)
    set_font(run2, size=11)

add_page_break(doc)


# ---------- 5. DSI ----------
add_h1(doc, "5. Doctrina Social de la Iglesia, ética y transparencia")

add_justified(doc,
    "La Doctrina Social de la Iglesia (DSI) propone principios éticos de larga tradición que "
    "resuenan con la ISO 26000 y enriquecen la fundamentación de nuestra propuesta. Seleccionamos "
    "dos principios particularmente pertinentes:")

add_h2(doc, "5.1 Bien común")
add_justified(doc,
    "La DSI entiende el bien común como el conjunto de condiciones sociales que permiten a las "
    "personas y comunidades alcanzar su perfeccionamiento (Compendio DSI, 164). El agua, los "
    "bosques, los acuíferos, los ecosistemas locales son bienes comunes por naturaleza: nadie "
    "puede apropiárselos para su beneficio particular sin lesionar a todos los demás.")
add_justified(doc,
    "El principio es directamente pertinente al problema que atendemos: la corrupción ambiental es "
    "una captura del bien común por intereses privados. La plataforma combate esa captura al "
    "distribuir la vigilancia entre muchos —la ciudadanía— y al publicar los patrones de "
    "irregularidad. Así ayuda a resolver la problemática: devuelve al bien común su carácter "
    "público y vigilable.")

add_h2(doc, "5.2 Subsidiariedad")
add_justified(doc,
    "El principio de subsidiariedad sostiene que las decisiones deben tomarse en el nivel más "
    "cercano posible a las personas afectadas, y que las instancias superiores deben apoyar —no "
    "sustituir— las capacidades de las menores (Compendio DSI, 185-188).")
add_justified(doc,
    "Este principio es clave en nuestro proyecto: la comunidad local es la primera en observar y "
    "reconocer la irregularidad; antes que delegar todo en la autoridad central, se le da un "
    "instrumento para actuar desde abajo. La autoridad nacional (PROFEPA, CONAGUA) no desaparece, "
    "pero recibe la información con la calidad, el detalle y la urgencia que sólo da la proximidad "
    "local. Esto ayuda a resolver la problemática porque corrige un error frecuente: centralizar "
    "la vigilancia ambiental en un Estado con recursos limitados.")

add_h2(doc, "5.3 Ética y transparencia operativa")
add_bullet(doc, "Declaración explícita de no discriminación ni persecución personal: se denuncian prácticas, no personas.")
add_bullet(doc, "Protección del denunciante: anonimato por defecto, cifrado de datos, no trazabilidad.")
add_bullet(doc, "Publicación trimestral de un tablero con números de denuncias, tipología y estatus.")
add_bullet(doc, "Auditoría anual externa por parte de la Universidad Anáhuac o de una ONG independiente.")

add_page_break(doc)


# ---------- 6. RECURSOS Y SOSTENIBILIDAD ----------
add_h1(doc, "6. Uso de recursos y sostenibilidad")

add_h2(doc, "6.1 Planificación de tiempo, espacio y recursos")
add_justified(doc,
    "El piloto se diseña a seis meses con tres fases, cada una con recursos explícitamente "
    "asignados y obstáculos anticipados:")

fases_tabla = [
    ["1", "Diseño", "2 meses", "1 coordinador, 1 programador, 1 asesor jurídico ambiental, 1 especialista educativo",
     "Costo alto de desarrollo custom → uso de Django + HTML/JS estándar y plantillas abiertas."],
    ["2", "Implementación piloto", "3 meses", "2 facilitadores + voluntarios universitarios",
     "Baja afluencia inicial → activación en una institución educativa con base de usuarios captiva."],
    ["3", "Evaluación y ajustes", "1 mes", "Equipo completo + voluntarios + presupuesto total $40-45K MXN",
     "Falta de seguimiento oficial → convenio previo con SEDESU/PROFEPA como receptores formales."],
]
add_table(doc,
    ["Fase", "Título", "Duración", "Recursos", "Obstáculo y superación"],
    fases_tabla,
    widths=[1.0, 2.8, 2.0, 4.2, 5.0])

add_h2(doc, "6.2 Impacto positivo a largo plazo")
add_h3(doc, "Ámbito social")
add_justified(doc,
    "Recuperación de confianza ciudadana en las instituciones, construcción de cultura de "
    "legalidad y fortalecimiento de capacidades comunitarias para la vigilancia colectiva. El "
    "impacto es acumulativo: cada denuncia exitosa aumenta la probabilidad de que otras "
    "personas reporten.")
add_h3(doc, "Ámbito ambiental")
add_justified(doc,
    "Reducción de la impunidad en tala ilegal, sobreexplotación hídrica y contaminación. A "
    "mediano plazo, mayor conservación de la Sierra Gorda y del acuífero del Valle de Querétaro. "
    "A largo plazo, una ciudadanía formada que protege ecosistemas de forma preventiva.")
add_h3(doc, "Ámbito económico")
add_justified(doc,
    "Menores costos públicos por remediación ambiental; competencia más leal entre empresas "
    "que sí cumplen la regulación; preservación de los servicios ecosistémicos (agua, suelo "
    "fértil) de los que depende el crecimiento económico de Querétaro.")

add_h2(doc, "6.3 Contribución a tres metas ODS")
ods_tabla = [
    ["16.5",
     "Reducir sustancialmente la corrupción y el soborno en todas sus formas",
     "Canal de denuncia ciudadana que eleva la rendición de cuentas (A en Klitgaard) y expone "
     "patrones de irregularidad ambiental."],
    ["13.3",
     "Educación, sensibilización y capacidad humana frente al cambio climático",
     "Seis micro-módulos educativos y talleres comunitarios que forman a la ciudadanía sobre "
     "impactos climáticos de la corrupción ambiental."],
    ["15.1",
     "Conservación y uso sostenible de los ecosistemas terrestres",
     "Denuncias específicas de tala ilegal y uso indebido del suelo; vínculo con PROFEPA para "
     "protección de áreas naturales locales."],
]
add_table(doc,
    ["Meta", "Descripción oficial", "Cómo contribuye Denuncia Verde"],
    ods_tabla,
    widths=[1.3, 5.5, 8.2])

add_page_break(doc)


# ---------- 7. DERECHOS HUMANOS ----------
add_h1(doc, "7. Derechos humanos")

add_justified(doc,
    "Nuestro proyecto se enmarca en tres derechos humanos directamente pertinentes. Para cada uno "
    "se describe su contenido, su fuente normativa y recomendaciones precisas de cómo abordarlo "
    "operativamente.")

add_h2(doc, "7.1 Derecho a un medio ambiente sano")
add_justified(doc,
    "Reconocido por la Asamblea General de la ONU en julio de 2022 (A/RES/76/300) y contenido en "
    "el artículo 4º de la Constitución Política de los Estados Unidos Mexicanos. Establece que toda "
    "persona tiene derecho a vivir en un entorno que permita su dignidad, salud y bienestar.")
add_bullet(doc, "Categorizar las denuncias según tipo de afectación ambiental para ofrecer evidencia estructurada.", bold_lead="Recomendación 1")
add_bullet(doc, "Firmar un convenio con PROFEPA para que cada denuncia canalizada reciba folio oficial y seguimiento trazable.", bold_lead="Recomendación 2")
add_bullet(doc, "Publicar semestralmente un indicador público de denuncias atendidas vs. denuncias resueltas.", bold_lead="Recomendación 3")

add_h2(doc, "7.2 Derecho a la participación ciudadana")
add_justified(doc,
    "Contenido en el artículo 21 de la Declaración Universal de los Derechos Humanos (DUDH, 1948) "
    "y desarrollado por el Acuerdo de Escazú (CEPAL, 2018) en materia ambiental. Garantiza que toda "
    "persona pueda participar en los asuntos públicos y en la toma de decisiones que le afectan.")
add_bullet(doc, "Habilitar denuncia 100% anónima para remover barreras de miedo a represalias.", bold_lead="Recomendación 1")
add_bullet(doc, "Organizar talleres trimestrales abiertos con organizaciones civiles y la Universidad Anáhuac.", bold_lead="Recomendación 2")
add_bullet(doc, "Traducir los módulos a lenguaje incluyente, accesible y adaptable a distintas comunidades.", bold_lead="Recomendación 3")

add_h2(doc, "7.3 Derecho al acceso a la información")
add_justified(doc,
    "Artículo 19 de la DUDH (1948) y artículo 6º de la Constitución mexicana. La persona tiene "
    "derecho a buscar, recibir y difundir información por cualquier medio. Incluye el acceso a la "
    "información pública ambiental (permisos, sanciones, inspecciones).")
add_bullet(doc, "Integrar en la plataforma enlaces directos a la Plataforma Nacional de Transparencia para consultar permisos ambientales.", bold_lead="Recomendación 1")
add_bullet(doc, "Mantener un glosario público de términos técnicos (vocabulario ambiental, legal, administrativo).", bold_lead="Recomendación 2")
add_bullet(doc, "Publicar todas las fuentes que sustenten los módulos educativos; ninguna afirmación sin referencia verificable.", bold_lead="Recomendación 3")

add_page_break(doc)


# ---------- 8. STAKEHOLDERS ----------
add_h1(doc, "8. Matriz de grupos de interés y reflexión sobre la transformación social")

add_h2(doc, "8.1 Matriz de grupos de interés")
stake_tabla = [
    ["Comunidad local",
     "Beneficiaria y usuaria directa",
     "Reporta irregularidades, asiste a talleres, vigila su entorno",
     "Recuperación de su entorno natural y fortalecimiento de su voz política"],
    ["Universidades (Anáhuac, UAQ)",
     "Aliadas estratégicas",
     "Investigación, voluntariado, asesoría técnica y jurídica",
     "Formación en proyectos reales y vinculación con el entorno"],
    ["Gobierno local (SEDESU) y federal (PROFEPA, CONAGUA)",
     "Receptores institucionales de denuncias canalizadas",
     "Seguimiento formal, inspección y sanción cuando procede",
     "Mejora de su reputación institucional y reducción de carga por filtrado ciudadano"],
    ["Organizaciones ambientales (Colectivo Pro Bosques, WWF local)",
     "Aliadas en contenidos y difusión",
     "Capacitación, validación técnica, amplificación en redes",
     "Escala de su trabajo y acceso a datos ciudadanos"],
    ["Empresas (agroindustria, construcción, extractivas)",
     "Sujeto regulado · aliado potencial",
     "Adoptar prácticas sostenibles, participar en mesas de transparencia",
     "Reducción de competencia desleal y reputación más sólida"],
    ["Medios de comunicación",
     "Amplificadores",
     "Difundir reportes y seguir casos relevantes",
     "Acceso a información pública estructurada y a fuentes verificables"],
]
add_table(doc,
    ["Grupo de interés", "Relación con el proyecto", "Contribución", "Impacto positivo"],
    stake_tabla,
    widths=[3.6, 3.4, 4.0, 4.0])

add_h2(doc, "8.2 Estrategias de colaboración institucional")
add_bullet(doc, "Convenio marco con PROFEPA y CONAGUA para la canalización oficial de denuncias con folio.", bold_lead="Gubernamental")
add_bullet(doc, "Vinculación con el Programa de Servicio Social de la Universidad Anáhuac para dotar al proyecto de voluntarios formados.", bold_lead="Académica")
add_bullet(doc, "Alianza con Transparencia Mexicana y Colectivo Pro Bosques para intercambiar datos y amplificar casos emblemáticos.", bold_lead="Sociedad civil")
add_bullet(doc, "Mesa semestral de diálogo con la Cámara Nacional de Empresarios de Querétaro para promover buenas prácticas.", bold_lead="Sector privado")

add_h2(doc, "8.3 Viabilidad desde las profesiones del equipo")
add_justified(doc,
    "Cada integrante aporta al proyecto un enfoque profesional distinto que, sumados, dan "
    "viabilidad real a la propuesta:")
perspectivas = [
    ("Juan Pablo Penedo — Actuaría",
     "Aporta el análisis cuantitativo: modelación de indicadores de impacto, tratamiento "
     "estadístico de las denuncias, identificación de patrones espacio-temporales de "
     "irregularidades y evaluación de riesgo por zona. Desde su formación, la transformación "
     "social se hace viable cuando la medimos con rigor y comunicamos su evolución con "
     "evidencia clara."),
    ("David Manuel Ruiz — Ingeniería Industrial",
     "Aporta la mirada de procesos: diseño del flujo operativo (recepción → validación → "
     "canalización → seguimiento), definición de indicadores de desempeño, mejora continua y "
     "eficiencia en el uso de los recursos del piloto. Desde su disciplina, la transformación "
     "social es viable cuando se traduce en sistemas repetibles, medibles y escalables."),
    ("Leonardo Mayoral — Administración y Finanzas",
     "Aporta la estrategia y la gestión: plan financiero del piloto, búsqueda de aliados y "
     "patrocinadores, diseño del modelo de sostenibilidad económica, administración de "
     "relaciones con stakeholders. Desde su profesión, la transformación social es viable "
     "cuando se construye un modelo organizacional robusto que sobreviva al piloto."),
    ("Pablo Andrés Pacheco — Derecho",
     "Aporta el fundamento normativo: adecuación a la Ley General del Equilibrio Ecológico, "
     "protección legal del denunciante, cumplimiento de la Ley Federal de Protección de Datos, "
     "articulación con la PROFEPA y las autoridades locales. Desde su disciplina, la "
     "transformación social es viable cuando cada acción ciudadana se ampara en un marco "
     "jurídico que la protege y le da efecto."),
]
for titulo, texto in perspectivas:
    add_h3(doc, titulo)
    add_justified(doc, texto)

add_page_break(doc)


# ---------- 9. PARTICIPACIÓN Y COLABORACIÓN ----------
add_h1(doc, "9. Participación y colaboración del equipo")

add_justified(doc,
    "El equipo trabajó de forma colaborativa y armónica durante todas las actividades "
    "parciales (1a, 1b, 2a, 3a, 3b, 3c) y en la integración final. Se adoptó una "
    "metodología de reuniones semanales, división clara de responsabilidades por bloques "
    "temáticos y revisión cruzada de todos los entregables. Las decisiones se tomaron por "
    "consenso, registrando los acuerdos en un documento compartido.")
add_justified(doc,
    "A continuación, cada integrante presenta su reflexión individual sobre su "
    "participación y su aprendizaje:")

reflexiones = [
    ("Juan Pablo Penedo Antúnez",
     "Actuaría",
     "Mi aporte principal fue el diseño de la estructura cuantitativa del proyecto: la "
     "manera de medir el impacto (número de denuncias por tipo, distribución geográfica, "
     "tasa de seguimiento) y la incorporación del modelo de Klitgaard como marco analítico. "
     "Coordiné además el prototipo técnico de la plataforma. Aprendí que la Responsabilidad "
     "Social exige combinar rigor analítico con sensibilidad humana: los números no "
     "sustituyen a las personas, las acompañan. Esta experiencia reforzó mi convicción de "
     "que las herramientas actuariales —medición, evaluación de riesgo, modelos de "
     "probabilidad— pueden ponerse al servicio de causas públicas, no sólo de empresas e "
     "instituciones financieras."),
    ("David Manuel Ruiz Pérez",
     "Ingeniería Industrial",
     "Mi contribución se centró en estructurar los procesos operativos del proyecto: el "
     "flujo de denuncia, el protocolo de validación, el cronograma de las tres fases del "
     "piloto y la medición de desempeño. También participé activamente en la redacción de "
     "los módulos educativos sobre cómo denunciar. Aprendí que un buen proyecto social "
     "necesita procesos tan claros como cualquier proyecto industrial: sin un flujo "
     "definido, la buena voluntad se pierde. También entendí que la sustentabilidad no es "
     "sólo ambiental, sino organizacional: un proyecto es sostenible si puede operar sin el "
     "equipo original que lo creó."),
    ("Leonardo Mayoral García",
     "Administración y Finanzas",
     "Mi rol fue coordinar la estrategia de alianzas, el plan financiero del piloto y la "
     "relación con los grupos de interés. Diseñé la matriz de stakeholders y el modelo de "
     "sostenibilidad económica a partir del piloto. Esta materia me mostró que la "
     "responsabilidad social no es un gasto ni un añadido de marketing, sino una inversión "
     "estratégica que genera valor compartido (Porter y Kramer) para la empresa, el "
     "gobierno y la sociedad. Me llevo una visión más amplia de lo que significa "
     "administrar: ya no sólo recursos financieros, sino confianza social y relaciones "
     "institucionales."),
    ("Pablo Andrés Pacheco Colmenares",
     "Derecho",
     "Mi contribución fue el análisis del marco jurídico: derechos humanos aplicables "
     "(ambiente sano, participación, información), alineación con la ISO 26000 y con la "
     "legislación mexicana, protección legal del denunciante y tratamiento de datos "
     "personales. Colaboré en la redacción del capítulo sobre DSI y en la articulación "
     "institucional con PROFEPA y CEDH. Aprendí que el Derecho no basta por sí solo: "
     "puede existir una ley impecable y al mismo tiempo una impunidad sistémica. Lo que "
     "cambia las cosas es sumar al marco jurídico una ciudadanía activa y organizada. "
     "Este proyecto me mostró, de forma concreta, cómo un abogado puede construir "
     "transformación social más allá del litigio."),
]
for nombre, carrera, reflexion in reflexiones:
    add_h3(doc, f"{nombre} — {carrera}")
    add_justified(doc, reflexion)

add_page_break(doc)


# ---------- 10. REFERENCIAS ----------
add_h1(doc, "10. Referencias")

refs = [
    "Carroll, A. B. (1991). The pyramid of corporate social responsibility: Toward the moral management of "
    "organizational stakeholders. Business Horizons, 34(4), 39-48.",
    "Comisión Económica para América Latina y el Caribe (CEPAL). (2018). Acuerdo regional sobre el "
    "acceso a la información, la participación pública y el acceso a la justicia en asuntos ambientales "
    "en América Latina y el Caribe (Acuerdo de Escazú). Naciones Unidas.",
    "Consejo Pontificio Justicia y Paz. (2005). Compendio de la Doctrina Social de la Iglesia. Libreria "
    "Editrice Vaticana.",
    "ISO. (2010). ISO 26000: Guía de Responsabilidad Social. International Organization for Standardization.",
    "Klitgaard, R. (1988). Controlling Corruption. University of California Press.",
    "Klitgaard, R. (1998). International cooperation against corruption. Finance and Development, 35(1), 3-6.",
    "Naciones Unidas. (1948). Declaración Universal de los Derechos Humanos. Asamblea General, París.",
    "Naciones Unidas. (2022). Resolución A/RES/76/300: Derecho humano a un medio ambiente limpio, "
    "saludable y sostenible. Asamblea General.",
    "Porter, M. E., & Kramer, M. R. (2011). Creating shared value. Harvard Business Review, 89(1/2), 62-77.",
    "PROFEPA. (2024). Informe anual de denuncias ambientales. Procuraduría Federal de Protección al "
    "Ambiente, México.",
    "Salamon, L. M. (2004). Global Civil Society: Dimensions of the Nonprofit Sector (Vol. 2). Kumarian "
    "Press.",
    "Sen, A. (1999). Development as Freedom. Alfred A. Knopf.",
    "Transparencia Mexicana. (2023). Índice de corrupción ambiental: estados con mayor incidencia. "
    "Ciudad de México.",
]
for r in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(r)
    set_font(run, size=11)


# ---------- Guardar ----------
DEST_DIR.mkdir(parents=True, exist_ok=True)
dest = DEST_DIR / FILENAME
doc.save(str(dest))
print(f"OK - Documento generado en:\n  {dest}")
